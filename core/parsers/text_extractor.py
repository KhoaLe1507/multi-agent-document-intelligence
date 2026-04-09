# core/parsers/text_extractor.py
import os
import tempfile
from pathlib import Path
import docx
from ..utils.logger import agent_logger
from ..exceptions.agent_errors import ParsingError
from ..schemas.data_types import DocumentChunk

def extract_text_chunks(file_path: str) -> list[DocumentChunk]:
    path = Path(file_path)
    chunks = []
    
    try:
        if path.suffix.lower() == '.docx':
            doc = docx.Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())
            
            # Đọc thêm table trong word (nếu có)
            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_data:
                        full_text.append(" | ".join(row_data))
                        
            content = "\n".join(full_text)
            chunks.append(DocumentChunk(
                chunk_type="text",
                content=content,
                file_name=path.name,
                page_number=1
            ))
            agent_logger.success(f"Đã trích xuất chữ từ file Word {path.name}.")
            
        elif path.suffix.lower() == '.doc':
            # Ghi chú: Định dạng .doc cũ (nhị phân) không được hỗ trợ chính thức bởi python-docx.
            # Chúng ta sẽ thử catch-all (nếu file thực chất là docx nhưng bị đối tên thành .doc)
            try:
                doc = docx.Document(file_path)
                full_text = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
                content = "\n".join(full_text)
                chunks.append(DocumentChunk(
                    chunk_type="text", content=content, file_name=path.name, page_number=1
                ))
                agent_logger.success(f"Đã đọc được {path.name} qua python-docx (do thực chất cấu trúc là XML).")
            except Exception:
                # Nếu thực sự là file binary OLE .doc cũ, ta dùng mẹo quét chuỗi printable text cơ bản
                agent_logger.warning(f"File {path.name} là dạng .doc cũ. Dùng fallback đọc raw text...")
                with open(file_path, "rb") as f:
                    content_bytes = f.read()
                    
                import string
                # Lọc ra các ký tự ASCII in được và một số ký tự Unicode cơ bản
                printable = set(string.printable.encode())
                extracted = []
                for b in content_bytes:
                    if b in printable:
                        extracted.append(chr(b))
                    elif b > 127: # Chấp nhận đại một số bytes utf-8 có thể
                        extracted.append(" ")
                        
                raw_string = "".join(extracted)
                # Dọn dẹp khoảng trắng
                words = [w for w in raw_string.split() if len(w) > 3]
                content = " ".join(words)
                
                if len(content) > 0:
                    chunks.append(DocumentChunk(
                        chunk_type="text", 
                        content=f"[WARNING: Đọc thô từ .doc binary, chữ có thể dính vào nhau hoặc mất dấu]\n{content}", 
                        file_name=path.name, 
                        page_number=1
                    ))
                else:
                    raise ValueError("Không tìm thấy text nào sau khi parse raw binary.")

        elif path.suffix.lower() == '.txt':
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            chunks.append(DocumentChunk(
                chunk_type="text", content=content, file_name=path.name, page_number=1
            ))
            agent_logger.success(f"Đã trích xuất text từ file {path.name}.")

        return chunks
        
    except Exception as e:
        agent_logger.error(f"Lỗi phân tích file văn bản {path.name}: {e}")
        raise ParsingError(f"Không thể parse file word/txt: {str(e)}")
